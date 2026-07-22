"""File handling utilities for resume uploads"""
import io
import aiofiles
import os
from pathlib import Path
from typing import Optional
from fastapi import UploadFile, HTTPException
import pdfplumber
from docx import Document
from app.core.config import settings
import logging

logger = logging.getLogger(__name__)

class FileHandler:
    """Handle file uploads, validation, and text extraction"""
    
    # Allowed file extensions
    ALLOWED_EXTENSIONS = getattr(settings, 'allowed_extensions', ['.pdf', '.docx', '.doc'])
    
    # Maximum file size (default 10MB)
    MAX_FILE_SIZE = getattr(settings, 'max_file_size', 10 * 1024 * 1024)
    
    @classmethod
    def validate_file(cls, file: UploadFile) -> None:
        """
        Validate file type and extension
        
        Args:
            file: UploadFile object
            
        Raises:
            HTTPException: If file type is not allowed
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in cls.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_ext} not allowed. Allowed types: {', '.join(cls.ALLOWED_EXTENSIONS)}"
            )
        
        logger.info(f"File validated: {file.filename} (extension: {file_ext})")
    
    @classmethod
    async def validate_file_size(cls, file: UploadFile) -> bytes:
        """
        Validate file size and read contents
        
        Args:
            file: UploadFile object
            
        Returns:
            bytes: File contents
            
        Raises:
            HTTPException: If file exceeds size limit
        """
        # Read file contents
        contents = await file.read()
        
        # Check file size
        if len(contents) > cls.MAX_FILE_SIZE:
            file_size_mb = len(contents) / (1024 * 1024)
            max_size_mb = cls.MAX_FILE_SIZE / (1024 * 1024)
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Size: {file_size_mb:.2f}MB, Max size: {max_size_mb:.0f}MB"
            )
        
        logger.info(f"File size validated: {len(contents)} bytes")
        return contents
    
    @classmethod
    def extract_text_from_pdf(cls, contents: bytes) -> str:
        """
        Extract text from PDF file using pdfplumber (layout-aware)
        
        Args:
            contents: PDF file bytes
            
        Returns:
            str: Extracted text
            
        Raises:
            HTTPException: If PDF extraction fails
        """
        try:
            text = ""
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                logger.info(f"Processing PDF with {len(pdf.pages)} pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Try layout-aware extraction first
                    page_text = page.extract_text(layout=True)
                    
                    # Fallback to simple text extraction if layout fails
                    if not page_text or not page_text.strip():
                        page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        text += f"\n--- Page {page_num} ---\n"
                        text += page_text + "\n\n"
                    
                    # Try to extract tables if needed
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:
                                    text += " | ".join(str(cell) for cell in row if cell) + "\n"
                            text += "\n"
            
            # Clean up extracted text
            text = cls._clean_extracted_text(text)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("PDF contains insufficient text content")
            
            logger.info(f"PDF extraction successful. Text length: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
    
    @classmethod
    def extract_text_from_docx(cls, contents: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            contents: DOCX file bytes
            
        Returns:
            str: Extracted text
            
        Raises:
            HTTPException: If DOCX extraction fails
        """
        try:
            doc = Document(io.BytesIO(contents))
            text = ""
            
            # Extract from paragraphs
            logger.info(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"
            
            # Extract headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
                
                # Footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
            
            # Clean up extracted text
            text = cls._clean_extracted_text(text)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("DOCX contains insufficient text content")
            
            logger.info(f"DOCX extraction successful. Text length: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
    
    @classmethod
    def extract_text_from_txt(cls, contents: bytes) -> str:
        """
        Extract text from TXT file
        
        Args:
            contents: TXT file bytes
            
        Returns:
            str: Extracted text
        """
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = contents.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # Fallback: ignore errors
                text = contents.decode('utf-8', errors='ignore')
            
            text = cls._clean_extracted_text(text)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("TXT file contains insufficient text content")
            
            logger.info(f"TXT extraction successful. Text length: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from TXT: {str(e)}")
    
    @classmethod
    def _clean_extracted_text(cls, text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and special characters
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove multiple newlines
        import re
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove excessive spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove empty lines at start and end
        text = text.strip()
        
        return text
    
    @classmethod
    async def save_upload_file(cls, contents: bytes, candidate_id: str, original_filename: str) -> str:
        """
        Save uploaded file to disk
        
        Args:
            contents: File bytes
            candidate_id: Candidate ID for filename
            original_filename: Original filename for extension
            
        Returns:
            str: Path to saved file
        """
        # Get file extension from original filename
        file_ext = Path(original_filename).suffix.lower()
        
        # Create safe filename using candidate ID
        safe_filename = f"{candidate_id}{file_ext}"
        file_path = settings.uploads_dir / safe_filename
        
        try:
            # Ensure upload directory exists
            settings.uploads_dir.mkdir(parents=True, exist_ok=True)
            
            # Save file asynchronously
            async with aiofiles.open(file_path, "wb") as f:
                await f.write(contents)
            
            logger.info(f"File saved successfully: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"File save error: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    @classmethod
    def extract_text(cls, contents: bytes, filename: str) -> str:
        """
        Extract text from file based on file extension
        
        Args:
            contents: File bytes
            filename: Original filename
            
        Returns:
            str: Extracted text
            
        Raises:
            HTTPException: If file type is unsupported or extraction fails
        """
        file_ext = Path(filename).suffix.lower()
        
        logger.info(f"Extracting text from: {filename} ({file_ext})")
        
        if file_ext == ".pdf":
            return cls.extract_text_from_pdf(contents)
        elif file_ext in [".docx", ".doc"]:
            return cls.extract_text_from_docx(contents)
        elif file_ext == ".txt":
            return cls.extract_text_from_txt(contents)
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_ext}. Supported: {', '.join(cls.ALLOWED_EXTENSIONS)}"
            )
    
    @classmethod
    def get_file_info(cls, filename: str, contents: bytes) -> dict:
        """
        Get file information without extracting text
        
        Args:
            filename: Original filename
            contents: File bytes
            
        Returns:
            dict: File information
        """
        file_ext = Path(filename).suffix.lower()
        file_size_mb = len(contents) / (1024 * 1024)
        
        return {
            "filename": filename,
            "extension": file_ext,
            "size_bytes": len(contents),
            "size_mb": round(file_size_mb, 2),
            "is_allowed": file_ext in cls.ALLOWED_EXTENSIONS,
            "max_size_mb": cls.MAX_FILE_SIZE / (1024 * 1024)
        }
    
    @classmethod
    async def delete_file(cls, file_path: str) -> bool:
        """
        Delete a file from disk
        
        Args:
            file_path: Path to file to delete
            
        Returns:
            bool: True if deleted, False otherwise
        """
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.info(f"File deleted: {file_path}")
                return True
            else:
                logger.warning(f"File not found for deletion: {file_path}")
                return False
        except Exception as e:
            logger.error(f"File deletion error: {e}")
            return False

# Convenience functions for easy import
def validate_file(file: UploadFile) -> None:
    """Convenience function for file validation"""
    FileHandler.validate_file(file)

async def extract_text_from_file(contents: bytes, filename: str) -> str:
    """Convenience function for text extraction"""
    return FileHandler.extract_text(contents, filename)

async def save_uploaded_file(contents: bytes, candidate_id: str, original_filename: str) -> str:
    """Convenience function for saving files"""
    return await FileHandler.save_upload_file(contents, candidate_id, original_filename)

def get_file_info(filename: str, contents: bytes) -> dict:
    """Convenience function for getting file info"""
    return FileHandler.get_file_info(filename, contents)